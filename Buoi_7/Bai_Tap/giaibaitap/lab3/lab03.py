import os, warnings
from pathlib import Path
warnings.filterwarnings('ignore')
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score
from sklearn.decomposition import PCA
from scipy.cluster.hierarchy import linkage, dendrogram

BASE=Path('.'); OUT=BASE/'lab03_outputs'; IMG=OUT/'images'
OUT.mkdir(exist_ok=True); IMG.mkdir(exist_ok=True)
PENG_PATH=BASE/'penguins.csv'; RETAIL_PATH=BASE/'OnlineRetail.csv'; RANDOM_STATE=42

def save_fig(path):
    plt.tight_layout(); plt.savefig(path, dpi=160, bbox_inches='tight'); plt.close(); return str(path)

def sil_score_fast(X, labels):
    n=len(labels); sample_size=min(1500,n) if n>1500 else None
    return float(silhouette_score(X, labels, sample_size=sample_size, random_state=RANDOM_STATE))

def metrics(X, labels):
    return {'Silhouette': sil_score_fast(X, labels),'Davies-Bouldin': float(davies_bouldin_score(X, labels)),'Calinski-Harabasz': float(calinski_harabasz_score(X, labels))}

def scan_kmeans(X, k_range=range(2,9), n_init=3):
    rows=[]
    for k in k_range:
        km=KMeans(n_clusters=k, init='k-means++', n_init=n_init, random_state=RANDOM_STATE)
        labels=km.fit_predict(X)
        rows.append({'k':k,'inertia':float(km.inertia_),'silhouette':sil_score_fast(X, labels)})
    return pd.DataFrame(rows)

def plot_elbow_sil(scan_df, title_prefix, path):
    fig, ax1=plt.subplots(figsize=(8,4.8))
    ax1.plot(scan_df['k'], scan_df['inertia'], marker='o', label='Inertia')
    ax1.set_xlabel('Số cụm K'); ax1.set_ylabel('Inertia / SSE'); ax1.grid(True, alpha=.25)
    ax2=ax1.twinx(); ax2.plot(scan_df['k'], scan_df['silhouette'], marker='s', linestyle='--', label='Silhouette'); ax2.set_ylabel('Silhouette Score')
    fig.suptitle(title_prefix + ' - Elbow & Silhouette')
    l1, lab1=ax1.get_legend_handles_labels(); l2, lab2=ax2.get_legend_handles_labels(); ax1.legend(l1+l2, lab1+lab2, loc='best')
    return save_fig(path)

def plot_pca_clusters(X_scaled, labels, title, path):
    pca=PCA(n_components=2, random_state=RANDOM_STATE); pts=pca.fit_transform(X_scaled)
    plt.figure(figsize=(8,4.8)); plt.scatter(pts[:,0], pts[:,1], c=labels, s=26, cmap='tab10', alpha=.85)
    plt.xlabel(f'PC1 ({pca.explained_variance_ratio_[0]*100:.1f}%)'); plt.ylabel(f'PC2 ({pca.explained_variance_ratio_[1]*100:.1f}%)')
    plt.title(title); plt.grid(True, alpha=.25); return save_fig(path)

def plot_dendrogram_sample(X_scaled, title, path, sample=160):
    rng=np.random.default_rng(RANDOM_STATE); idx=rng.choice(np.arange(X_scaled.shape[0]), size=min(sample, X_scaled.shape[0]), replace=False)
    Z=linkage(X_scaled[idx], method='ward', metric='euclidean')
    plt.figure(figsize=(9,4.8)); dendrogram(Z, truncate_mode='lastp', p=30, leaf_rotation=90., show_contracted=True)
    plt.title(title); plt.xlabel('Nhóm / kích thước cụm'); plt.ylabel('Khoảng cách'); return save_fig(path)

def profile_table(df, labels, cols):
    tmp=df.copy(); tmp['Cluster']=labels
    return tmp.groupby('Cluster')[cols].agg(['count','mean','median','min','max']).round(2)

# Penguins
peng_raw=pd.read_csv(PENG_PATH); peng=peng_raw.dropna().copy()
if 'sex' in peng.columns:
    peng=peng[peng['sex'].isin(['MALE','FEMALE'])].copy(); peng['sex_num']=peng['sex'].map({'FEMALE':0,'MALE':1})
num_cols=[c for c in ['culmen_length_mm','culmen_depth_mm','flipper_length_mm','body_mass_g','sex_num'] if c in peng.columns]
Xp=peng[num_cols].copy(); Xp_scaled=StandardScaler().fit_transform(Xp)
scan_p=scan_kmeans(Xp_scaled, range(2,9)); best_k_p=int(scan_p.loc[scan_p['silhouette'].idxmax(),'k']); k_p=3
labels_km_p=KMeans(n_clusters=k_p, init='k-means++', n_init=3, random_state=RANDOM_STATE).fit_predict(Xp_scaled)
labels_ag_p=AgglomerativeClustering(n_clusters=k_p, linkage='ward').fit_predict(Xp_scaled)
metrics_p_km=metrics(Xp_scaled, labels_km_p); metrics_p_ag=metrics(Xp_scaled, labels_ag_p)
peng_scan_path=plot_elbow_sil(scan_p, 'Penguins', IMG/'penguins_elbow_silhouette.png')
peng_km_path=plot_pca_clusters(Xp_scaled, labels_km_p, 'Penguins - K-Means clustering (K=3)', IMG/'penguins_kmeans_pca.png')
peng_ag_path=plot_pca_clusters(Xp_scaled, labels_ag_p, 'Penguins - Agglomerative clustering (K=3)', IMG/'penguins_agglomerative_pca.png')
peng_den_path=plot_dendrogram_sample(Xp_scaled, 'Penguins - Dendrogram mẫu', IMG/'penguins_dendrogram.png')

# Online retail
retail0=pd.read_csv(RETAIL_PATH, encoding='latin1')
retail0['InvoiceDate']=pd.to_datetime(retail0['InvoiceDate'], errors='coerce')
retail=retail0.dropna(subset=['CustomerID','InvoiceDate'])
retail=retail[(retail['Quantity']>0)&(retail['UnitPrice']>0)].copy()
retail['CustomerID']=retail['CustomerID'].astype(int).astype(str); retail['TotalPrice']=retail['Quantity']*retail['UnitPrice']
snapshot=retail['InvoiceDate'].max()+pd.Timedelta(days=1)
rfm=retail.groupby('CustomerID').agg(Recency=('InvoiceDate', lambda x: (snapshot-x.max()).days),Frequency=('InvoiceNo','nunique'),Monetary=('TotalPrice','sum'),Quantity=('Quantity','sum'),AvgBasket=('TotalPrice','mean'),Country=('Country', lambda x: x.mode().iat[0] if not x.mode().empty else 'Unknown')).reset_index()
rfm_model=rfm.copy()
for c in ['Recency','Frequency','Monetary','Quantity','AvgBasket']:
    lo,hi=rfm_model[c].quantile([0.01,0.99]); rfm_model[c]=rfm_model[c].clip(lo,hi)
rfm_log=rfm_model.copy()
for c in ['Recency','Frequency','Monetary','Quantity','AvgBasket']:
    rfm_log[c]=np.log1p(rfm_log[c])
retail_cols=['Recency','Frequency','Monetary','Quantity','AvgBasket']
Xr_scaled=StandardScaler().fit_transform(rfm_log[retail_cols])
scan_r=scan_kmeans(Xr_scaled, range(2,9), n_init=3); best_k_r=int(scan_r.loc[scan_r['silhouette'].idxmax(),'k']); k_r=4
labels_km_r=KMeans(n_clusters=k_r, init='k-means++', n_init=3, random_state=RANDOM_STATE).fit_predict(Xr_scaled)
metrics_r_km=metrics(Xr_scaled, labels_km_r)
# Agglomerative on a deterministic sample to avoid O(n^2) blow-up on full retail data
rng=np.random.default_rng(RANDOM_STATE); sample_idx=rng.choice(np.arange(Xr_scaled.shape[0]), size=min(1200, Xr_scaled.shape[0]), replace=False)
Xr_sample=Xr_scaled[sample_idx]
labels_ag_r=AgglomerativeClustering(n_clusters=k_r, linkage='ward').fit_predict(Xr_sample)
metrics_r_ag=metrics(Xr_sample, labels_ag_r)
rfm_labeled=rfm.copy(); rfm_labeled['Cluster']=labels_km_r
cent=rfm_labeled.groupby('Cluster')[retail_cols].mean(); score=cent['Monetary'].rank()+cent['Frequency'].rank()+(-cent['Recency']).rank(); order=score.sort_values(ascending=False).index.tolist()
role_candidates=['VIP / Giá trị cao','Trung thành tiềm năng','Ngủ quên / cần kích hoạt lại','Mới hoặc giá trị thấp']; role_names={cl:(role_candidates[i] if i<len(role_candidates) else f'Nhóm {cl}') for i,cl in enumerate(order)}
rfm_labeled['SegmentName']=rfm_labeled['Cluster'].map(role_names)
cluster_summary=rfm_labeled.groupby(['Cluster','SegmentName']).agg(Customers=('CustomerID','count'),Recency=('Recency','mean'),Frequency=('Frequency','mean'),Monetary=('Monetary','mean'),Quantity=('Quantity','mean'),AvgBasket=('AvgBasket','mean')).reset_index().round(2)
retail_scan_path=plot_elbow_sil(scan_r, 'Online Retail RFM', IMG/'retail_elbow_silhouette.png')
retail_km_path=plot_pca_clusters(Xr_scaled, labels_km_r, 'Online Retail - K-Means RFM clustering (K=4)', IMG/'retail_kmeans_pca.png')
retail_ag_path=plot_pca_clusters(Xr_sample, labels_ag_r, 'Online Retail - Agglomerative RFM clustering mẫu (K=4)', IMG/'retail_agglomerative_pca.png')
retail_den_path=plot_dendrogram_sample(Xr_scaled, 'Online Retail RFM - Dendrogram mẫu', IMG/'retail_dendrogram.png', sample=180)
plt.figure(figsize=(9,4.8)); plt.bar(cluster_summary['SegmentName'], cluster_summary['Monetary']); plt.title('Online Retail - Doanh thu trung bình theo cụm'); plt.ylabel('Monetary trung bình'); plt.xticks(rotation=20, ha='right'); retail_bar_path=save_fig(IMG/'retail_cluster_monetary.png')

# Save CSV
scan_p.to_csv(OUT/'penguins_k_scan.csv', index=False); scan_r.to_csv(OUT/'retail_k_scan.csv', index=False); cluster_summary.to_csv(OUT/'retail_cluster_summary.csv', index=False); rfm_labeled.to_csv(OUT/'retail_rfm_clustered.csv', index=False)
metrics_df=pd.DataFrame([{'Dataset':'Penguins','Model':'K-Means K=3', **metrics_p_km},{'Dataset':'Penguins','Model':'Agglomerative K=3', **metrics_p_ag},{'Dataset':'Online Retail RFM','Model':'K-Means K=4', **metrics_r_km},{'Dataset':'Online Retail RFM','Model':'Agglomerative K=4 (sample)', **metrics_r_ag}]).round(4)
metrics_df.to_csv(OUT/'lab03_metrics_summary.csv', index=False)

# DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_table(doc, df, max_rows=20):
    df=df.copy(); df=df.head(max_rows)
    table=doc.add_table(rows=1, cols=len(df.columns)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    for j,c in enumerate(df.columns): table.rows[0].cells[j].text=str(c); shade(table.rows[0].cells[j],'D9EAF7')
    for _,row in df.iterrows():
        cells=table.add_row().cells
        for j,c in enumerate(df.columns):
            val=row[c]; cells[j].text=(f'{val:,.4f}' if isinstance(val,float) and abs(val)<10 else f'{val:,.2f}' if isinstance(val,float) else str(val))
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    return table

def add_img(doc, path, caption, width=6.1):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path), width=Inches(width))
    cap=doc.add_paragraph(caption); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.runs[0].italic=True

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72)
for s in ['Normal','Heading 1','Heading 2','Heading 3']:
    doc.styles[s].font.name='Arial'; doc.styles[s]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
doc.styles['Normal'].font.size=Pt(10.5)
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=t.add_run('BÁO CÁO LAB03\nCÁC GIẢI THUẬT PHÂN CỤM CƠ BẢN'); r.bold=True; r.font.size=Pt(18)
doc.add_paragraph('Nội dung: K-Means, phân cụm đa cấp, đánh giá mức độ ứng dụng').alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.add_run('Tóm tắt: ').bold=True; p.add_run('Báo cáo triển khai phân cụm trên Penguins và Online Retail. Penguins dùng để minh họa phân nhóm theo đặc trưng hình thái; Online Retail được chuyển thành RFM để phân nhóm khách hàng phục vụ marketing và CRM.')

doc.add_heading('1. Mục tiêu và phương pháp', level=1)
doc.add_paragraph('Lab03 tập trung vào K-Means và phân cụm đa cấp. Các bước chính gồm tiền xử lý dữ liệu, chuẩn hóa, thử nhiều số cụm K, đánh giá bằng Elbow/Silhouette và diễn giải kết quả theo khả năng ứng dụng thực tế.')
doc.add_paragraph('Vì phân cụm là học không giám sát, báo cáo không dùng accuracy mà dùng Silhouette, Davies-Bouldin, Calinski-Harabasz và phân tích ý nghĩa cụm.')

doc.add_heading('2. Dataset Penguins - Phân cụm chim cánh cụt', level=1)
doc.add_paragraph(f'Dữ liệu ban đầu có {peng_raw.shape[0]} dòng, sau làm sạch còn {peng.shape[0]} dòng. Đặc trưng dùng cho mô hình: {", ".join(num_cols)}. Dataset hiện tại không có cột species nên báo cáo chỉ đánh giá cấu trúc cụm, không đối chiếu nhãn loài thật.')
doc.add_paragraph(f'K tốt nhất theo Silhouette khi quét 2-8 là {best_k_p}. Tuy nhiên báo cáo chọn K=3 làm cấu hình chính để bám sát bài toán phân cụm loài chim cánh cụt.')
add_table(doc, scan_p.round(4)); add_img(doc,peng_scan_path,'Hình 1. Penguins - Elbow và Silhouette theo K')
add_img(doc,peng_km_path,'Hình 2. Penguins - K-Means K=3 sau PCA'); add_img(doc,peng_ag_path,'Hình 3. Penguins - Agglomerative K=3 sau PCA'); add_img(doc,peng_den_path,'Hình 4. Penguins - Dendrogram mẫu')
doc.add_heading('2.1. Đánh giá và mức độ ứng dụng', level=2); add_table(doc, metrics_df[metrics_df['Dataset']=='Penguins'])
doc.add_paragraph('K-Means và Agglomerative có Silhouette gần nhau khi K=3. Điều này cho thấy dữ liệu có cấu trúc nhóm tương đối nhưng chưa tách biệt hoàn toàn. Mức độ ứng dụng phù hợp cho EDA sinh học, kiểm tra các cá thể có tự gom thành nhóm hình thái hay không. Muốn đánh giá chuẩn hơn cần thêm nhãn species để so sánh sau phân cụm.')

doc.add_heading('3. Dataset Online Retail - Phân cụm khách hàng bằng RFM', level=1)
doc.add_paragraph(f'Sau làm sạch giao dịch, dữ liệu tạo được bảng RFM gồm {rfm.shape[0]} khách hàng. Các biến dùng gồm Recency, Frequency, Monetary, Quantity và AvgBasket. Dữ liệu được xử lý ngoại lai, log-transform và chuẩn hóa trước khi phân cụm.')
doc.add_paragraph(f'K tốt nhất theo Silhouette khi quét 2-8 là {best_k_r}. Báo cáo chọn K=4 để tạo nhóm khách hàng đủ dễ hiểu và dễ triển khai trong marketing.')
add_table(doc, scan_r.round(4)); add_img(doc,retail_scan_path,'Hình 5. Online Retail - Elbow và Silhouette theo K')
add_img(doc,retail_km_path,'Hình 6. Online Retail - K-Means RFM K=4 sau PCA'); add_img(doc,retail_ag_path,'Hình 7. Online Retail - Agglomerative trên mẫu RFM K=4 sau PCA'); add_img(doc,retail_den_path,'Hình 8. Online Retail - Dendrogram mẫu'); add_img(doc,retail_bar_path,'Hình 9. Monetary trung bình theo cụm')
doc.add_heading('3.1. Hồ sơ cụm khách hàng', level=2); add_table(doc, cluster_summary)
doc.add_paragraph('Nhóm VIP / giá trị cao có Monetary và Frequency cao, Recency thấp, nên ưu tiên giữ chân. Nhóm trung thành tiềm năng có thể upsell/cross-sell. Nhóm ngủ quên cần chiến dịch kích hoạt lại. Nhóm mới hoặc giá trị thấp cần onboarding, voucher nhỏ và gợi ý sản phẩm phù hợp.')
doc.add_heading('3.2. Đánh giá và mức độ ứng dụng', level=2); add_table(doc, metrics_df[metrics_df['Dataset']=='Online Retail RFM'])
doc.add_paragraph('Online Retail là phần có tính ứng dụng cao nhất. Cụm khách hàng có thể chuyển thành hành động trong CRM: cá nhân hóa ưu đãi, chăm sóc VIP, phát hiện khách hàng có nguy cơ rời bỏ, tối ưu ngân sách marketing và xây dashboard phân khúc khách hàng.')

doc.add_heading('4. So sánh K-Means và phân cụm đa cấp', level=1)
compare=pd.DataFrame({'Tiêu chí':['Cách hoạt động','Điểm mạnh','Hạn chế','Khi nên dùng'],'K-Means':['Chọn K, khởi tạo centroid, gán điểm vào cụm gần nhất và cập nhật lặp lại','Nhanh, dễ mở rộng, phù hợp dữ liệu lớn','Cần chọn K trước, nhạy outlier và thang đo','Khi cần phân nhóm nhanh để ứng dụng CRM/marketing'],'Agglomerative':['Mỗi điểm là một cụm ban đầu rồi gộp dần theo khoảng cách/linkage','Có dendrogram, dễ giải thích cấu trúc phân cấp','Tốn tài nguyên hơn với dữ liệu lớn','Khi muốn phân tích quan hệ phân cấp hoặc dữ liệu không quá lớn']})
add_table(doc, compare)
doc.add_heading('5. Kết luận', level=1)
doc.add_paragraph('Lab03 cho thấy phân cụm là công cụ mạnh để khám phá cấu trúc ẩn. Penguins giúp minh họa phân cụm theo hình thái, còn Online Retail có khả năng ứng dụng trực tiếp vào kinh doanh. Khi triển khai thực tế, cần chuẩn hóa dữ liệu, xử lý ngoại lai, thử nhiều K và chọn số cụm không chỉ theo điểm số mà còn theo khả năng diễn giải và hành động.')
doc.add_heading('Phụ lục: Cách chạy code', level=1)
doc.add_paragraph('Đặt file lab03_super_report.py cùng thư mục với penguins.csv và OnlineRetail.csv, sau đó chạy:')
q=doc.add_paragraph(); q.style='Intense Quote'; q.add_run('pip install pandas numpy matplotlib scikit-learn scipy python-docx\npython lab03_super_report.py')
docx_path=OUT/'Lab03_BaoCao_PhanCum.docx'; doc.save(docx_path)
# standalone code path uses relative BASE
src=Path(__file__).read_text(encoding='utf-8').replace("BASE=Path('.'); OUT=BASE/'lab03_outputs'; IMG=OUT/'images'", "BASE=Path('.'); OUT=BASE/'lab03_outputs'; IMG=OUT/'images'")
(BASE/'lab03_super_report.py').write_text(src, encoding='utf-8')
print('DONE', docx_path, BASE/'lab03_super_report.py')
