# -*- coding: utf-8 -*-
"""
LAB04 - Tập phổ biến và luật kết hợp
Tác vụ: Apriori + Association Rules trên dữ liệu Online Retail.

Cách chạy:
    pip install pandas numpy matplotlib python-docx
    python lab04.py

Đầu ra:
    lab04_outputs/
        Lab04_BaoCao_Apriori_LuatKetHop.docx
        frequent_itemsets.csv
        association_rules.csv
        *.png

Ghi chú:
- Code dùng thuật toán Apriori tự cài đặt, không bắt buộc mlxtend/apyori.
- Nếu máy yếu, giảm MAX_TRANSACTIONS xuống 3000.
"""

from __future__ import annotations

from collections import Counter
from itertools import combinations
from pathlib import Path
import math
import warnings

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

warnings.filterwarnings("ignore")

# =========================
# 1. Cấu hình
# =========================
DATA_PATHS = [
    "OnlineRetail.csv",
    "Online Retail.csv",
    "Online Retail.xlsx",
    "online_retail.csv",
]
OUT_DIR = Path("lab04_outputs")
IMG_DIR = OUT_DIR / "images"

# Giới hạn để chạy nhanh nhưng vẫn đủ kết quả như bài lab.
MAX_TRANSACTIONS = 5000
MIN_SUPPORT = 0.02          # 2% tổng số hóa đơn
MIN_CONFIDENCE = 0.60       # 60%
MAX_ITEMSET_LEN = 3         # 3 là vừa đủ đẹp và không quá chậm
TOP_RULES = 20

RANDOM_STATE = 42
np.random.seed(RANDOM_STATE)


# =========================
# 2. Tiện ích đọc dữ liệu
# =========================
def find_data_file() -> Path:
    """Tìm file dữ liệu Online Retail trong thư mục hiện tại."""
    for p in DATA_PATHS:
        path = Path(p)
        if path.exists():
            return path
    # Tìm rộng hơn theo tên gần đúng
    for path in Path(".").glob("*Retail*"):
        if path.suffix.lower() in [".csv", ".xlsx", ".xls"]:
            return path
    raise FileNotFoundError(
        "Không tìm thấy OnlineRetail.csv / Online Retail.xlsx trong thư mục chạy. "
        "Hãy đặt file dữ liệu cùng thư mục với lab04.py."
    )


def read_online_retail(path: Path) -> pd.DataFrame:
    """Đọc CSV/XLSX linh hoạt."""
    if path.suffix.lower() in [".xlsx", ".xls"]:
        return pd.read_excel(path)
    # thử nhiều encoding vì dataset này hay lỗi ký tự
    for enc in ["utf-8", "latin1", "ISO-8859-1"]:
        try:
            return pd.read_csv(path, encoding=enc)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(path, encoding_errors="ignore")


def clean_online_retail(df: pd.DataFrame) -> pd.DataFrame:
    """Làm sạch dữ liệu giao dịch: bỏ hóa đơn hủy, số lượng âm, mô tả rỗng."""
    required = {"InvoiceNo", "Description", "Quantity"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Thiếu cột bắt buộc: {missing}. File hiện có: {df.columns.tolist()}")

    clean = df.copy()
    clean["InvoiceNo"] = clean["InvoiceNo"].astype(str)
    clean["Description"] = clean["Description"].astype(str).str.strip().str.upper()

    clean = clean[
        (clean["Quantity"] > 0)
        & (~clean["InvoiceNo"].str.startswith("C"))
        & (clean["Description"].notna())
        & (clean["Description"] != "")
        & (clean["Description"].str.lower() != "nan")
    ].copy()

    # Loại các dòng mô tả lỗi phổ biến nếu có
    bad_words = ["POSTAGE", "DOTCOM POSTAGE", "MANUAL", "AMAZON FEE", "BANK CHARGES"]
    clean = clean[~clean["Description"].isin(bad_words)]
    return clean


def build_transactions(df: pd.DataFrame, max_transactions: int = MAX_TRANSACTIONS) -> list[set[str]]:
    """Chuyển dữ liệu nhiều dòng/hóa đơn thành danh sách giao dịch."""
    invoice_ids = df["InvoiceNo"].drop_duplicates().tolist()
    if len(invoice_ids) > max_transactions:
        # Lấy mẫu có random_state để kết quả tái lập
        invoice_ids = pd.Series(invoice_ids).sample(max_transactions, random_state=RANDOM_STATE).tolist()
    sub = df[df["InvoiceNo"].isin(invoice_ids)]

    transactions = []
    for _, g in sub.groupby("InvoiceNo"):
        items = set(g["Description"].dropna().astype(str).str.strip().str.upper())
        if len(items) >= 2:
            transactions.append(items)
    return transactions


# =========================
# 3. Apriori tự cài đặt
# =========================
def apriori(transactions: list[set[str]], min_support: float, max_len: int = 3):
    """
    Tìm tập phổ biến bằng Apriori.
    Trả về:
        frequent_counts: dict[frozenset, int]
        frequent_by_k: dict[int, dict[frozenset, int]]
    """
    n_transactions = len(transactions)
    min_count = math.ceil(min_support * n_transactions)

    # L1
    item_counter = Counter()
    for trans in transactions:
        item_counter.update(trans)

    L1 = {
        frozenset([item]): count
        for item, count in item_counter.items()
        if count >= min_count
    }

    frequent_by_k = {1: L1}
    frequent_counts = dict(L1)

    prev_frequent = set(L1.keys())

    for k in range(2, max_len + 1):
        if not prev_frequent:
            break

        # Sinh candidate bằng cách đếm tổ hợp k-item trong từng giao dịch,
        # nhưng chỉ dùng các item đã phổ biến ở bước 1 để giảm nhiễu.
        frequent_items = set()
        for itemset in prev_frequent:
            frequent_items.update(itemset)

        candidate_counter = Counter()
        for trans in transactions:
            filtered = sorted(trans & frequent_items)
            if len(filtered) < k:
                continue
            for comb in combinations(filtered, k):
                comb_set = frozenset(comb)
                # Apriori prune: mọi subset k-1 phải phổ biến
                valid = all(frozenset(s) in prev_frequent for s in combinations(comb_set, k - 1))
                if valid:
                    candidate_counter[comb_set] += 1

        Lk = {
            itemset: count
            for itemset, count in candidate_counter.items()
            if count >= min_count
        }
        if not Lk:
            break

        frequent_by_k[k] = Lk
        frequent_counts.update(Lk)
        prev_frequent = set(Lk.keys())

    return frequent_counts, frequent_by_k


def generate_rules(frequent_counts: dict[frozenset, int], n_transactions: int, min_confidence: float) -> pd.DataFrame:
    """Sinh luật kết hợp từ các tập phổ biến."""
    rows = []
    for itemset, itemset_count in frequent_counts.items():
        if len(itemset) < 2:
            continue

        items = list(itemset)
        for r in range(1, len(items)):
            for antecedent_tuple in combinations(items, r):
                antecedent = frozenset(antecedent_tuple)
                consequent = itemset - antecedent

                antecedent_count = frequent_counts.get(antecedent)
                consequent_count = frequent_counts.get(consequent)
                if not antecedent_count or not consequent_count:
                    continue

                support = itemset_count / n_transactions
                antecedent_support = antecedent_count / n_transactions
                consequent_support = consequent_count / n_transactions
                confidence = support / antecedent_support if antecedent_support else 0
                lift = confidence / consequent_support if consequent_support else 0
                leverage = support - antecedent_support * consequent_support
                conviction = (
                    (1 - consequent_support) / (1 - confidence)
                    if confidence < 1 else np.inf
                )

                if confidence >= min_confidence:
                    rows.append({
                        "antecedents": ", ".join(sorted(antecedent)),
                        "consequents": ", ".join(sorted(consequent)),
                        "antecedent_support": antecedent_support,
                        "consequent_support": consequent_support,
                        "support": support,
                        "confidence": confidence,
                        "lift": lift,
                        "leverage": leverage,
                        "conviction": conviction,
                        "itemset_len": len(itemset),
                    })

    if not rows:
        return pd.DataFrame(columns=[
            "antecedents", "consequents", "antecedent_support", "consequent_support",
            "support", "confidence", "lift", "leverage", "conviction", "itemset_len"
        ])

    rules = pd.DataFrame(rows)
    rules = rules.sort_values(["lift", "confidence", "support"], ascending=False).reset_index(drop=True)
    return rules


def frequent_itemsets_to_df(frequent_counts: dict[frozenset, int], n_transactions: int) -> pd.DataFrame:
    rows = []
    for itemset, count in frequent_counts.items():
        rows.append({
            "itemsets": ", ".join(sorted(itemset)),
            "length": len(itemset),
            "count": count,
            "support": count / n_transactions,
        })
    df = pd.DataFrame(rows).sort_values(["length", "support"], ascending=[True, False]).reset_index(drop=True)
    return df


# =========================
# 4. Biểu đồ
# =========================
def save_top_items_chart(frequent_df: pd.DataFrame, path: Path):
    top = frequent_df[frequent_df["length"] == 1].head(15).copy()
    if top.empty:
        return None
    plt.figure(figsize=(10, 6))
    plt.barh(top["itemsets"][::-1], top["support"][::-1])
    plt.xlabel("Support")
    plt.title("Top 15 sản phẩm phổ biến nhất")
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


def save_rules_scatter(rules: pd.DataFrame, path: Path):
    if rules.empty:
        return None
    plt.figure(figsize=(8, 6))
    sizes = np.clip(rules["lift"].replace(np.inf, np.nan).fillna(rules["lift"].replace(np.inf, np.nan).max()).values, 1, 20) * 18
    plt.scatter(rules["support"], rules["confidence"], s=sizes, alpha=0.55)
    plt.xlabel("Support")
    plt.ylabel("Confidence")
    plt.title("Mối quan hệ giữa Support và Confidence của luật kết hợp")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


def save_confidence_hist(rules: pd.DataFrame, path: Path):
    if rules.empty:
        return None
    plt.figure(figsize=(8, 5))
    plt.hist(rules["confidence"], bins=20)
    plt.xlabel("Confidence")
    plt.ylabel("Số lượng luật")
    plt.title("Phân phối độ tin cậy của các luật kết hợp")
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


def save_lift_bar(rules: pd.DataFrame, path: Path):
    if rules.empty:
        return None
    top = rules.head(10).copy()
    labels = (top["antecedents"] + " -> " + top["consequents"]).str.slice(0, 80)
    plt.figure(figsize=(10, 7))
    plt.barh(labels[::-1], top["lift"][::-1])
    plt.xlabel("Lift")
    plt.title("Top 10 luật có Lift cao nhất")
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return path


# =========================
# 5. Báo cáo DOCX
# =========================
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_df_table(doc, df: pd.DataFrame, max_rows=10, float_cols=None):
    if df.empty:
        doc.add_paragraph("Không có dữ liệu để hiển thị.")
        return
    show = df.head(max_rows).copy()
    if float_cols is None:
        float_cols = show.select_dtypes(include=[float, np.float64]).columns.tolist()
    for c in float_cols:
        if c in show.columns:
            show[c] = show[c].apply(lambda x: "∞" if x == np.inf else f"{x:.4f}")
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(show.columns):
        hdr[i].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_docx_report(
    output_path: Path,
    data_path: Path,
    raw_shape,
    clean_shape,
    n_transactions: int,
    frequent_df: pd.DataFrame,
    rules: pd.DataFrame,
    images: list[Path],
):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("LAB04 - TẬP PHỔ BIẾN VÀ LUẬT KẾT HỢP", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Bài thực hành Apriori trên dữ liệu giao dịch mua hàng tại siêu thị / Online Retail")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Mục tiêu bài thực hành", 1)
    doc.add_paragraph(
        "Bài thực hành triển khai thuật toán Apriori nhằm tìm các tập sản phẩm thường xuyên xuất hiện "
        "cùng nhau trong hóa đơn mua hàng. Từ các tập phổ biến, chương trình trích xuất luật kết hợp "
        "và đánh giá luật bằng các chỉ số support, confidence, lift, leverage và conviction."
    )

    add_heading(doc, "2. Cơ sở lý thuyết tóm tắt", 1)
    theory = [
        ("Support", "Tỷ lệ giao dịch chứa tập sản phẩm đang xét. Support càng cao chứng tỏ tập sản phẩm xuất hiện càng phổ biến."),
        ("Confidence", "Xác suất khách mua sản phẩm vế phải khi đã mua sản phẩm vế trái. Đây là chỉ số quan trọng để đánh giá độ tin cậy của luật."),
        ("Lift", "Đo mức tăng xác suất mua chung so với trường hợp độc lập. Lift > 1 cho thấy hai nhóm sản phẩm có xu hướng liên quan tích cực."),
        ("Leverage", "Mức chênh lệch giữa xác suất mua chung thực tế và xác suất kỳ vọng nếu hai nhóm sản phẩm độc lập."),
        ("Conviction", "Đo mức độ phụ thuộc một chiều của luật. Conviction càng lớn thì luật càng đáng chú ý."),
    ]
    for name, desc in theory:
        doc.add_paragraph(f"{name}: {desc}", style=None)

    add_heading(doc, "3. Dữ liệu và tiền xử lý", 1)
    doc.add_paragraph(f"File dữ liệu sử dụng: {data_path.name}")
    doc.add_paragraph(f"Kích thước dữ liệu gốc: {raw_shape[0]:,} dòng x {raw_shape[1]} cột")
    doc.add_paragraph(f"Kích thước sau làm sạch: {clean_shape[0]:,} dòng x {clean_shape[1]} cột")
    doc.add_paragraph(f"Số hóa đơn/giao dịch đưa vào Apriori: {n_transactions:,}")
    doc.add_paragraph(
        "Các bước tiền xử lý gồm: loại hóa đơn hủy bắt đầu bằng ký tự C, loại dòng có Quantity <= 0, "
        "loại mô tả sản phẩm rỗng, chuẩn hóa tên sản phẩm thành chữ hoa và gom nhiều dòng sản phẩm "
        "thành một giao dịch theo InvoiceNo."
    )

    add_heading(doc, "4. Kết quả tìm tập phổ biến", 1)
    doc.add_paragraph(
        f"Ngưỡng min_support = {MIN_SUPPORT:.2%}. Thuật toán tìm được {len(frequent_df):,} tập phổ biến. "
        "Bảng dưới đây trình bày các tập phổ biến có support cao nhất."
    )
    add_df_table(doc, frequent_df[["itemsets", "length", "count", "support"]], max_rows=12)

    add_heading(doc, "5. Kết quả luật kết hợp", 1)
    doc.add_paragraph(
        f"Ngưỡng min_confidence = {MIN_CONFIDENCE:.0%}. Chương trình sinh được {len(rules):,} luật kết hợp. "
        "Các luật được sắp xếp ưu tiên theo lift, confidence và support."
    )
    if not rules.empty:
        display_cols = ["antecedents", "consequents", "support", "confidence", "lift", "leverage", "conviction"]
        add_df_table(doc, rules[display_cols], max_rows=10)
    else:
        doc.add_paragraph("Không có luật nào thỏa min_confidence. Có thể giảm MIN_SUPPORT hoặc MIN_CONFIDENCE trong code.")

    add_heading(doc, "6. Trực quan hóa kết quả", 1)
    captions = {
        "top_items.png": "Hình 1. Top sản phẩm phổ biến nhất theo support.",
        "rules_scatter.png": "Hình 2. Biểu đồ quan hệ giữa support và confidence của các luật.",
        "confidence_hist.png": "Hình 3. Phân phối confidence của luật kết hợp.",
        "lift_bar.png": "Hình 4. Top luật có lift cao nhất.",
    }
    for img in images:
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(6.1))
            cap = doc.add_paragraph(captions.get(img.name, img.name))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "7. Đánh giá mức độ ứng dụng", 1)
    doc.add_paragraph(
        "Kết quả Apriori có thể ứng dụng trực tiếp vào bài toán Market Basket Analysis. "
        "Các luật có confidence cao cho biết khi khách hàng mua nhóm sản phẩm A thì có xác suất cao mua thêm sản phẩm B. "
        "Doanh nghiệp có thể dùng thông tin này để thiết kế combo sản phẩm, gợi ý bán kèm, bố trí hàng hóa gần nhau, "
        "tạo chương trình khuyến mãi chéo và cá nhân hóa đề xuất sản phẩm trên website thương mại điện tử."
    )
    doc.add_paragraph(
        "Tuy nhiên, cần lưu ý rằng luật kết hợp không khẳng định quan hệ nhân quả. Một luật có confidence cao nhưng lift thấp "
        "có thể chỉ phản ánh sản phẩm vế phải vốn đã rất phổ biến. Vì vậy, khi chọn luật để ứng dụng thực tế nên ưu tiên đồng thời "
        "support đủ lớn, confidence cao và lift > 1. Ngoài ra, dữ liệu cần được làm sạch tốt để tránh nhiễu từ hóa đơn hủy, sản phẩm lỗi hoặc mô tả không nhất quán."
    )

    add_heading(doc, "8. Kết luận", 1)
    doc.add_paragraph(
        "Bài thực hành đã hoàn thành hai nội dung chính: xác định tập phổ biến bằng Apriori và trích xuất luật kết hợp từ tập phổ biến. "
        "Thông qua các chỉ số support, confidence, lift, leverage và conviction, người phân tích có thể chọn ra những luật có giá trị thực tiễn "
        "để hỗ trợ quyết định kinh doanh trong bán lẻ và thương mại điện tử."
    )

    doc.save(output_path)


# =========================
# 6. Main
# =========================
def main():
    OUT_DIR.mkdir(exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    print("Đang tìm file dữ liệu Online Retail...")
    data_path = find_data_file()
    print(f"Tìm thấy: {data_path}")

    print("Đang đọc dữ liệu...")
    raw = read_online_retail(data_path)
    raw_shape = raw.shape

    print("Đang làm sạch dữ liệu...")
    clean = clean_online_retail(raw)
    clean_shape = clean.shape

    print("Đang gom hóa đơn thành giao dịch...")
    transactions = build_transactions(clean, MAX_TRANSACTIONS)
    n_transactions = len(transactions)
    print(f"Số giao dịch dùng để phân tích: {n_transactions:,}")

    if n_transactions == 0:
        raise ValueError("Không có giao dịch hợp lệ sau tiền xử lý.")

    print("Đang chạy Apriori...")
    frequent_counts, frequent_by_k = apriori(transactions, MIN_SUPPORT, MAX_ITEMSET_LEN)
    frequent_df = frequent_itemsets_to_df(frequent_counts, n_transactions)
    frequent_csv = OUT_DIR / "frequent_itemsets.csv"
    frequent_df.to_csv(frequent_csv, index=False, encoding="utf-8-sig")

    print("Đang sinh luật kết hợp...")
    rules = generate_rules(frequent_counts, n_transactions, MIN_CONFIDENCE)
    rules_csv = OUT_DIR / "association_rules.csv"
    rules.to_csv(rules_csv, index=False, encoding="utf-8-sig")

    # Xuất thống kê theo độ dài tập phổ biến
    summary_rows = []
    for k, d in frequent_by_k.items():
        summary_rows.append({"itemset_length": k, "num_frequent_itemsets": len(d)})
    pd.DataFrame(summary_rows).to_csv(OUT_DIR / "apriori_summary.csv", index=False, encoding="utf-8-sig")

    print("Đang vẽ biểu đồ...")
    images = []
    images.append(save_top_items_chart(frequent_df, IMG_DIR / "top_items.png"))
    images.append(save_rules_scatter(rules, IMG_DIR / "rules_scatter.png"))
    images.append(save_confidence_hist(rules, IMG_DIR / "confidence_hist.png"))
    images.append(save_lift_bar(rules, IMG_DIR / "lift_bar.png"))
    images = [p for p in images if p is not None]

    print("Đang tạo báo cáo Word...")
    docx_path = OUT_DIR / "Lab04_BaoCao_Apriori_LuatKetHop.docx"
    build_docx_report(
        output_path=docx_path,
        data_path=data_path,
        raw_shape=raw_shape,
        clean_shape=clean_shape,
        n_transactions=n_transactions,
        frequent_df=frequent_df,
        rules=rules,
        images=images,
    )

    print("\nHOÀN TẤT LAB04")
    print(f"- Báo cáo Word: {docx_path}")
    print(f"- Tập phổ biến: {frequent_csv}")
    print(f"- Luật kết hợp: {rules_csv}")
    print("\nTop 5 luật kết hợp:")
    if not rules.empty:
        print(rules[["antecedents", "consequents", "support", "confidence", "lift"]].head(5).to_string(index=False))
    else:
        print("Không có luật thỏa ngưỡng.")


if __name__ == "__main__":
    main()
